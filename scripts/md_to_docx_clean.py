"""
Clean markdown to Word conversion - optimized for non-technical users.
Removes excess whitespace and formats clearly for commissioners.
"""
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import sys

def markdown_to_docx_clean(md_file_path, docx_file_path):
    """Convert markdown to Word with clean, readable formatting for non-technical users."""
    try:
        # Read markdown file
        with open(md_file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Create Word document
        doc = Document()
        
        # Set default font and style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Set paragraph spacing - minimal for readability
        paragraph_format = style.paragraph_format
        paragraph_format.space_after = Pt(3)
        paragraph_format.space_before = Pt(0)
        
        lines = md_content.split('\n')
        i = 0
        prev_was_empty = False
        
        while i < len(lines):
            line = lines[i].rstrip()
            
            # Skip multiple consecutive empty lines
            if not line:
                if not prev_was_empty:
                    prev_was_empty = True
                i += 1
                continue
            
            prev_was_empty = False
            
            # Headers with appropriate spacing
            if line.startswith('# '):
                if len(doc.paragraphs) > 0:  # Add space before heading if not first element
                    doc.add_paragraph()
                heading = doc.add_heading(line[2:].strip(), level=1)
                heading_format = heading.paragraph_format
                heading_format.space_after = Pt(6)
                heading_format.space_before = Pt(0)
                i += 1
            elif line.startswith('## '):
                if len(doc.paragraphs) > 0:
                    doc.add_paragraph()
                heading = doc.add_heading(line[3:].strip(), level=2)
                heading_format = heading.paragraph_format
                heading_format.space_after = Pt(4)
                heading_format.space_before = Pt(0)
                i += 1
            elif line.startswith('### '):
                heading = doc.add_heading(line[4:].strip(), level=3)
                heading_format = heading.paragraph_format
                heading_format.space_after = Pt(3)
                heading_format.space_before = Pt(0)
                i += 1
            elif line.startswith('#### '):
                heading = doc.add_heading(line[5:].strip(), level=4)
                heading_format = heading.paragraph_format
                heading_format.space_after = Pt(3)
                heading_format.space_before = Pt(0)
                i += 1
            # Horizontal rule - convert to simple line
            elif line.strip().startswith('---') and len(line.strip()) <= 5:
                if len(doc.paragraphs) > 0:
                    doc.add_paragraph()
                i += 1
            # Code blocks - format as regular text (no code formatting for non-technical users)
            elif line.startswith('```'):
                i += 1
                # Skip code block content
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    i += 1
                if i < len(lines):
                    i += 1
            # Bullet points
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                text = line.strip()[2:].strip()
                if text:  # Only add if there's actual content
                    p = doc.add_paragraph(style='List Bullet')
                    # Handle bold in bullet points
                    if '**' in text:
                        parts = re.split(r'(\*\*.*?\*\*)', text)
                        for part in parts:
                            if part.startswith('**') and part.endswith('**'):
                                run = p.add_run(part[2:-2])
                                run.bold = True
                            else:
                                p.add_run(part)
                    else:
                        p.add_run(text)
                i += 1
            # Numbered lists
            elif re.match(r'^\d+\.\s', line.strip()):
                text = re.sub(r'^\d+\.\s', '', line.strip()).strip()
                if text:
                    p = doc.add_paragraph(style='List Number')
                    if '**' in text:
                        parts = re.split(r'(\*\*.*?\*\*)', text)
                        for part in parts:
                            if part.startswith('**') and part.endswith('**'):
                                run = p.add_run(part[2:-2])
                                run.bold = True
                            else:
                                p.add_run(part)
                    else:
                        p.add_run(text)
                i += 1
            # Checkbox lists - convert to regular bullet
            elif line.strip().startswith('- [') or line.strip().startswith('* ['):
                text = re.sub(r'^[-*]\s*\[[x\s]\]\s*', '', line.strip()).strip()
                if text:
                    p = doc.add_paragraph(style='List Bullet')
                    if '**' in text:
                        parts = re.split(r'(\*\*.*?\*\*)', text)
                        for part in parts:
                            if part.startswith('**') and part.endswith('**'):
                                run = p.add_run(part[2:-2])
                                run.bold = True
                            else:
                                p.add_run(part)
                    else:
                        p.add_run(text)
                i += 1
            # Regular paragraph
            else:
                # Skip lines that are just formatting characters
                if line.strip() in ['---', '===']:
                    i += 1
                    continue
                
                p = doc.add_paragraph()
                # Handle inline formatting (bold, italic)
                # Split on bold markers
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                        run = p.add_run(part[1:-1])
                        run.italic = True
                    elif part.startswith('`') and part.endswith('`'):
                        # Code inline - just remove backticks for non-technical users
                        p.add_run(part[1:-1])
                    else:
                        p.add_run(part)
                i += 1
        
        # Remove any trailing empty paragraphs
        while len(doc.paragraphs) > 0 and not doc.paragraphs[-1].text.strip():
            doc._body._body.remove(doc.paragraphs[-1]._p)
        
        # Save document
        docx_file_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(docx_file_path))
        
        file_size = docx_file_path.stat().st_size
        print(f"[OK] Converted: {md_file_path.name} -> {docx_file_path.name} ({file_size:,} bytes)")
        return True
        
    except Exception as e:
        print(f"[ERROR] Error converting {md_file_path.name}: {str(e)}")
        return False

def convert_framework_docs_clean():
    """Convert all framework documentation MD files to clean DOCX format."""
    script_dir = Path(__file__).parent
    
    DOCUMENTATION_FILES = [
        "COMMISSIONER_INSTRUCTIONS.md",
        "README.md",
        "FRAMEWORK_OVERVIEW.md",
        "DISTRIBUTION_CHECKLIST.md",
        "FILES_CHECKLIST.md"
    ]
    
    print("Converting framework documentation files to clean Word format...")
    print("="*60)
    
    converted_count = 0
    failed_count = 0
    
    for md_filename in DOCUMENTATION_FILES:
        md_path = script_dir / md_filename
        
        if not md_path.exists():
            print(f"[SKIP] {md_filename} - File not found")
            continue
        
        docx_filename = md_filename.replace('.md', '.docx')
        docx_path = script_dir / docx_filename
        
        print(f"\nConverting: {md_filename}")
        success = markdown_to_docx_clean(md_path, docx_path)
        
        if success:
            converted_count += 1
        else:
            failed_count += 1
    
    print("\n" + "="*60)
    print(f"Conversion complete: {converted_count} succeeded, {failed_count} failed")
    print("="*60)

if __name__ == "__main__":
    convert_framework_docs_clean()
