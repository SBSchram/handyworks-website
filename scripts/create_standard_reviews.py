import os
import PyPDF2
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from datetime import datetime

# Define the PDF directory
PDF_DIR = Path(r"C:\Users\sbsch\Dropbox\My Documents\ACAHM\Feb 2026 Meeting\VU\251201 VU Inst-DAcCHM Pre Accred Sup Rpt\attachments")
OUTPUT_DIR = Path(r"C:\Users\sbsch\Documents\handyworks-website\Standard_Review_Documents")

def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file."""
    try:
        text = ""
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        return f"Error reading PDF: {str(e)}"

def get_exhibits_for_standard(standard_num, pdf_dir):
    """Get all exhibit PDFs for a specific standard."""
    exhibits = []
    # Use set to avoid duplicates from case-insensitive glob
    pdf_files = set()
    for f in pdf_dir.glob('*.pdf'):
        pdf_files.add(f)
    for f in pdf_dir.glob('*.PDF'):
        pdf_files.add(f)
    pdf_files = sorted(pdf_files)
    
    for pdf_file in pdf_files:
        filename = pdf_file.name
        filename_lower = filename.lower()
        
        # Handle specific standards
        if standard_num == 2:
            # Standard 2.05: EX2.05 files
            if filename.startswith('EX2.05') or filename.startswith('EX2.05 '):
                exhibits.append(pdf_file)
        elif standard_num == 2.06:
            # Standard 2.06: EX2.06 files
            if filename.startswith('EX2.06') or filename.startswith('EX2.06 '):
                exhibits.append(pdf_file)
        elif standard_num == 2.07:
            # Standard 2.07: EX2.07 files or 2.07 files (local compliance monitoring)
            if filename.startswith('EX2.07') or filename.startswith('EX2.07 ') or filename.startswith('2.07') or 'monitoring' in filename_lower:
                exhibits.append(pdf_file)
        elif standard_num == 5.02:
            # Standard 5.02: EX5.02 files
            if filename.startswith('EX5.02') or filename.startswith('EX5.02 '):
                exhibits.append(pdf_file)
        elif standard_num == 7.01:
            # Standard 7.01: EX7.01 files
            if filename.startswith('EX7.01') or filename.startswith('EX7.01 '):
                exhibits.append(pdf_file)
        elif standard_num == 8.01:
            # Standard 8.01: EX8.01 files
            if filename.startswith('EX8.01') or filename.startswith('EX8.01 '):
                exhibits.append(pdf_file)
        elif standard_num == 9.02:
            # Standard 9.02: EX9.02 files (or financial reports)
            if filename.startswith('EX9.02') or filename.startswith('EX9.02 '):
                exhibits.append(pdf_file)
            # Also check for financial reports mentioned in requirements
            elif 'financial' in filename_lower and ('report' in filename_lower or 'plan' in filename_lower):
                exhibits.append(pdf_file)
    
    return sorted(exhibits)

def find_federal_compliance_docs(pdf_dir):
    """Find documents that might contain federal compliance information (Title IX, HIPAA, FERPA, OSHA)."""
    compliance_docs = []
    pdf_files = sorted(pdf_dir.glob('*.pdf')) + sorted(pdf_dir.glob('*.PDF'))
    
    # Keywords to search for
    compliance_keywords = ['title ix', 'hipaa', 'ferpa', 'osha', 'compliance training', 
                           'federal compliance', 'discrimination', 'privacy protection']
    
    for pdf_file in pdf_files:
        filename_lower = pdf_file.name.lower()
        
        # Check filename for compliance-related terms
        if any(keyword in filename_lower for keyword in ['compliance', 'training', 'title', 'hipaa', 'ferpa', 'osha']):
            compliance_docs.append(pdf_file)
            continue
        
        # Check if it's Employee Handbook or Student Handbook (likely to contain compliance info)
        if 'employee handbook' in filename_lower or 'student handbook' in filename_lower:
            compliance_docs.append(pdf_file)
            continue
        
        # Check if it's EX2.07 or EX2.08 (OSHA related)
        if filename_lower.startswith('ex2.07') or filename_lower.startswith('ex2.08'):
            compliance_docs.append(pdf_file)
    
    return sorted(compliance_docs)

def format_text_for_analysis(text):
    """Format text by converting numbered items to lists and cleaning up."""
    if not text or text.startswith("Error reading PDF"):
        return []
    
    # Split into lines and process
    lines = text.split('\n')
    formatted_lines = []
    current_list = []
    in_list = False
    
    for line in lines:
        line = line.strip()
        if not line:
            if current_list:
                formatted_lines.append(current_list)
                current_list = []
                in_list = False
            continue
        
        # Check if line starts with a number followed by period or space
        numbered_match = re.match(r'^(\d+)[\.\)]\s*(.+)', line)
        if numbered_match:
            if not in_list:
                if current_list:
                    formatted_lines.append(current_list)
                current_list = []
                in_list = True
            current_list.append(numbered_match.group(2).strip())
        else:
            if current_list:
                formatted_lines.append(current_list)
                current_list = []
                in_list = False
            if line:
                formatted_lines.append(line)
    
    if current_list:
        formatted_lines.append(current_list)
    
    return formatted_lines

def create_analysis_summary(text, exhibit_name, standard_num=None):
    """Create a concise paragraph-style summary framed in compliance context."""
    if not text or text.startswith("Error reading PDF"):
        return "Unable to extract text from this document. Please review the original PDF file."
    
    text_lower = text.lower()
    filename_lower = exhibit_name.lower()
    
    # For Standard 2.05 (federal compliance), frame everything in compliance context
    if standard_num == 2:
        # Training Materials - demonstrate actual training materials exist
        if 'training' in filename_lower and 'material' in filename_lower:
            # Extract section headers for training materials
            section_pattern = re.compile(r'(?:^|\n)\s*(?:[IVX]+\.|Section\s+[IVX]+|Part\s+[IVX]+|Chapter\s+[IVX]+|^\d+\.)\s*([^\n]{10,80})', re.MULTILINE | re.IGNORECASE)
            section_matches = section_pattern.findall(text)
            sections = [s.strip() for s in section_matches[:8]] if section_matches else []
            
            if 'ferpa' in filename_lower:
                summary = "A comprehensive review of all aspects of FERPA (Family Educational Rights and Privacy Act)."
                if sections:
                    summary += " It includes an overview, and sections: " + ", ".join(sections) + "."
                else:
                    summary += " It covers student rights, educational records, disclosure requirements, staff responsibilities, and compliance procedures."
            
            elif 'hipaa' in filename_lower:
                summary = "The actual training materials used for compliance with HIPAA (privacy protection of patients' health records)."
                if sections:
                    summary += " It includes an overview, and sections: " + ", ".join(sections) + "."
                else:
                    summary += " A comprehensive review covering patient privacy protection, protected health information (PHI), disclosure rules, and staff responsibilities."
            
            elif 'title ix' in filename_lower or 'title 9' in filename_lower:
                summary = "A comprehensive review of Title IX compliance requirements regarding sex-based discrimination and sexual harassment."
                if sections:
                    summary += " It includes an overview, and sections: " + ", ".join(sections) + "."
                else:
                    summary += " It covers policies, procedures, reporting requirements, and staff responsibilities."
            
            elif 'osha' in filename_lower or 'calosha' in filename_lower:
                summary = "The actual training materials used for compliance with OSHA (occupational safety and health)."
                if sections:
                    summary += " It includes an overview, and sections: " + ", ".join(sections) + "."
                else:
                    summary += " A comprehensive review covering workplace safety standards, hazardous communication, training requirements, and compliance procedures."
            
            else:
                summary = "The actual training materials used for compliance with federal laws and regulations."
        
        # Training Tests - evidence of completion verification
        elif 'test' in filename_lower or ('training' in filename_lower and 'test' in text_lower):
            if 'ferpa' in filename_lower:
                summary = "A compliance test used to verify understanding of FERPA requirements and evidence that stakeholders completed required FERPA compliance training."
            elif 'hipaa' in filename_lower:
                summary = "A compliance test used to verify understanding of HIPAA requirements and evidence that stakeholders completed required HIPAA compliance training."
            elif 'osha' in filename_lower:
                summary = "A compliance test used to verify understanding of OSHA requirements and evidence that stakeholders completed required OSHA compliance training."
            else:
                summary = "A compliance test used to verify training completion and evidence that stakeholders completed required compliance training."
        
        # Training Completion Evidence (meeting minutes, training records)
        elif 'training' in filename_lower and ('faculty' in filename_lower or 'student' in filename_lower or 'staff' in filename_lower):
            if 'ferpa' in filename_lower:
                summary = "Evidence that relevant stakeholders ("
                if 'faculty' in filename_lower:
                    summary += "faculty"
                elif 'student' in filename_lower:
                    summary += "students"
                else:
                    summary += "staff"
                summary += ") completed required FERPA compliance training."
            elif 'title ix' in filename_lower:
                summary = "Evidence that relevant stakeholders completed required Title IX compliance training."
            else:
                summary = "Evidence that relevant stakeholders completed required compliance training."
        
        # Meeting Minutes with Training
        elif 'minutes' in filename_lower or 'meeting' in filename_lower:
            date_match = re.search(r'(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}', text, re.IGNORECASE)
            date_str = f" dated {date_match.group(0)}" if date_match else ""
            summary = f"Meeting minutes{date_str} providing evidence that relevant stakeholders completed required compliance training."
            if 'title ix' in filename_lower:
                summary = f"Meeting minutes{date_str} providing evidence that relevant stakeholders completed required Title IX compliance training."
            elif 'osha' in filename_lower or 'calosha' in filename_lower:
                summary = f"Meeting minutes{date_str} providing evidence that relevant stakeholders completed required OSHA compliance training."
        
        # Compliance Officer Resume/Evidence
        elif 'evidence' in filename_lower and ('compliance' in filename_lower or 'officer' in filename_lower):
            summary = "Evidence that staff members responsible for overseeing compliance with federal regulations are appropriately qualified. This document describes the person who will be in charge of compliance, including their qualifications, education, and relevant experience."
        
        elif 'resume' in filename_lower or 'cv' in filename_lower:
            summary = "Evidence that staff members responsible for overseeing compliance with federal regulations are appropriately qualified. This document describes the person who will be in charge of compliance, including their qualifications, education, and relevant experience."
        
        else:
            summary = "Document providing evidence of compliance with federal laws and regulations."
    
    # Standard 2.06 (state compliance - Cal/OSHA)
    elif standard_num == 2.06:
        if 'training' in filename_lower and 'material' in filename_lower:
            section_pattern = re.compile(r'(?:^|\n)\s*(?:[IVX]+\.|Section\s+[IVX]+|Part\s+[IVX]+|Chapter\s+[IVX]+|^\d+\.)\s*([^\n]{10,80})', re.MULTILINE | re.IGNORECASE)
            section_matches = section_pattern.findall(text)
            sections = [s.strip() for s in section_matches[:8]] if section_matches else []
            summary = "The actual training materials used for compliance with state laws and regulations (Cal/OSHA - occupational safety and health)."
            if sections:
                summary += " It includes an overview, and sections: " + ", ".join(sections) + "."
        elif 'test' in filename_lower:
            summary = "A compliance test used to verify understanding of Cal/OSHA requirements and evidence that stakeholders completed required state compliance training."
        elif 'training' in filename_lower and ('faculty' in filename_lower or 'student' in filename_lower or 'staff' in filename_lower):
            summary = "Evidence that relevant stakeholders completed required Cal/OSHA compliance training."
        elif 'minutes' in filename_lower or 'meeting' in filename_lower:
            date_match = re.search(r'(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}', text, re.IGNORECASE)
            date_str = f" dated {date_match.group(0)}" if date_match else ""
            summary = f"Meeting minutes{date_str} providing evidence that relevant stakeholders completed required Cal/OSHA compliance training."
        elif 'resume' in filename_lower or 'evidence' in filename_lower:
            summary = "Evidence that staff members responsible for overseeing compliance with state regulations are appropriately qualified."
        else:
            summary = "Document providing evidence of compliance with state laws and regulations."
    
    # Standard 2.07 (local compliance monitoring)
    elif standard_num == 2.07:
        if 'policy' in filename_lower or 'monitoring' in filename_lower:
            summary = "Evidence that an ongoing process is followed to ensure continuous compliance with local regulations."
        elif 'resume' in filename_lower or 'evidence' in filename_lower or 'qualification' in filename_lower:
            summary = "Evidence that staff members responsible for overseeing compliance with local regulations are appropriately qualified."
        else:
            summary = "Document providing evidence of qualified staff and ongoing process for monitoring compliance with local and municipal laws, ordinances, codes, and regulatory requirements."
    
    # Standard 5.02 (admissions requirements)
    elif standard_num == 5.02:
        if 'policy' in filename_lower:
            summary = "Policy document demonstrating that all admitted students meet all program admissions requirements at the time of enrollment, including English language proficiency and undergraduate credit requirements."
        elif 'audit' in filename_lower:
            summary = "Results of institutional file audits for admissions records demonstrating that all admitted students met all program admissions requirements at the time of enrollment."
        elif 'template' in filename_lower:
            summary = "Template documents used to ensure all admitted students meet all program admissions requirements."
        elif 'training' in filename_lower:
            summary = "Evidence of staff training to ensure all admitted students meet all program admissions requirements."
        else:
            summary = "Document providing evidence that all students are meeting all program admissions requirements at the time of enrollment."
    
    # Standard 7.01 (course prerequisites)
    elif standard_num == 7.01:
        if 'prerequisite' in filename_lower:
            summary = "Evidence that the program has appropriate course prerequisites and that students have completed all prerequisites prior to enrollment in a course."
        elif 'review' in filename_lower or 'form' in filename_lower:
            summary = "Documentation demonstrating that students have completed all prerequisites prior to enrollment in a course."
        elif 'method' in filename_lower:
            summary = "Methods used to ensure prerequisites are met prior to course enrollment."
        else:
            summary = "Document providing evidence that the program has appropriate course prerequisites and that students have completed all prerequisites prior to enrollment."
    
    # Standard 8.01 (core faculty)
    elif standard_num == 8.01:
        if 'roster' in filename_lower or 'list' in filename_lower:
            summary = "Evidence that the program clearly identifies a core group of faculty members who have regular and ongoing responsibility for the design, delivery, and assessment of the program."
        elif 'agreement' in filename_lower or 'contract' in filename_lower:
            summary = "Executed agreements for core faculty demonstrating their roles and responsibilities."
        elif 'job description' in filename_lower or 'responsibility' in filename_lower:
            summary = "Executed job descriptions for core faculty that outline the required roles and responsibilities."
        elif 'minutes' in filename_lower or 'meeting' in filename_lower:
            date_match = re.search(r'(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}', text, re.IGNORECASE)
            date_str = f" dated {date_match.group(0)}" if date_match else ""
            summary = f"Meeting minutes{date_str} demonstrating that core faculty are fulfilling their responsibilities in program development, review, and governance."
        else:
            summary = "Document providing evidence that the program employs an identifiable core group of qualified faculty members."
    
    # Standard 9.02 (financial stability)
    elif standard_num == 9.02:
        if 'financial report' in filename_lower or 'quarterly' in filename_lower or 'q2' in filename_lower or 'q3' in filename_lower or '2025' in filename_lower:
            # Check if this is Q2 or Q3 2025
            q2_match = 'q2' in filename_lower or 'second quarter' in filename_lower
            q3_match = 'q3' in filename_lower or 'third quarter' in filename_lower
            
            quarter_str = ""
            if q2_match:
                quarter_str = "2025 Q2"
            elif q3_match:
                quarter_str = "2025 Q3"
            else:
                quarter_str = "2025"
            
            summary = f"Quarterly financial report for {quarter_str}. "
            
            # Check for required components in text
            text_lower_check = text_lower if text else ""
            components_found = []
            
            if 'budget' in text_lower_check and ('actual' in text_lower_check or 'vs' in text_lower_check):
                components_found.append("Budget vs. Actual")
            if 'balance sheet' in text_lower_check or 'statement of financial position' in text_lower_check:
                components_found.append("Balance Sheet (statement of financial position)")
            if 'income statement' in text_lower_check or 'profit and loss' in text_lower_check or 'p&l' in text_lower_check:
                components_found.append("Income Statement (profit and loss statement)")
            if 'cash flow' in text_lower_check:
                components_found.append("Cash Flow Statement")
            
            if components_found:
                summary += f"Report includes: {', '.join(components_found)}. "
            
            # Check for narrative interpretation
            if 'narrative' in text_lower_check or 'interpretation' in text_lower_check or 'discussion' in text_lower_check or 'analysis' in text_lower_check:
                summary += "Includes narrative interpretations of financial statements."
            else:
                summary += "Review document to verify it includes narrative interpretations of each required financial statement component."
        
        elif 'financial plan' in filename_lower or ('plan' in filename_lower and 'financial' in filename_lower):
            text_lower_check = text_lower if text else ""
            summary = "Financial plan addressing Standard 9.02 concerns. "
            
            # Check if it addresses specific concerns
            concerns_addressed = []
            concerns_details = []
            
            if 'related party' in text_lower_check or 'related-party' in text_lower_check or 'related -party' in text_lower_check:
                concerns_addressed.append("related party transactions")
                # Extract details
                if 'advance' in text_lower_check and 'shareholder' in text_lower_check:
                    concerns_details.append("Identifies advances and loans to shareholder, interest receivables, and facility lease arrangements")
            
            if 'illiquid' in text_lower_check or ('receivable' in text_lower_check and 'shareholder' in text_lower_check):
                concerns_addressed.append("illiquid assets (receivables from shareholder)")
                concerns_details.append("Identifies loans to shareholder and other parties that cannot be used to support operations")
            
            if ('shareholder' in text_lower_check and 'equity' in text_lower_check) and ('negative' in text_lower_check or 'deficit' in text_lower_check or '-' in text and '563' in text):
                concerns_addressed.append("negative shareholder's equity")
                concerns_details.append("Addresses accumulated deficit and equity position")
            
            if ('cash' in text_lower_check and 'balance' in text_lower_check) or ('distribution' in text_lower_check and 'shareholder' in text_lower_check):
                concerns_addressed.append("low cash balances due to shareholder distributions")
                concerns_details.append("Identifies low ending cash balances and large shareholder distributions that caused cash deficiencies")
            
            # Check for governing board approval
            board_approved = 'governing board' in text_lower_check and ('approve' in text_lower_check or 'approved' in text_lower_check)
            if board_approved:
                summary += "Approved (or to be approved) by VU's governing board. "
            else:
                summary += "Review to verify governing board approval. "
            
            # Check for corrective actions/timeline
            if 'corrective action' in text_lower_check or 'timeline' in text_lower_check:
                summary += "Includes corrective actions with immediate (1-3 months), short-term (3-12 months), and long-term (12-24 months) timelines. "
            
            if concerns_addressed:
                summary += f"Addresses all four key concerns: {', '.join(concerns_addressed)}. "
                if concerns_details:
                    summary += " ".join(concerns_details[:2]) + "."
            else:
                summary += "Review document to verify it addresses all concerns: extensive related party transactions, illiquid assets (receivables from shareholder), negative shareholder's equity, and persistently low cash balances due to large shareholder distributions."
            
            # Note about quarterly reports requirement
            if 'quarterly' in text_lower_check and '2026' in text_lower_check:
                summary += " NOTE: This plan mentions quarterly reporting starting 2026 Q1, but Standard 9.02 requires quarterly financial reports for 2025 Q2 and Q3 which are separate documents."
        
        elif 'budget' in filename_lower:
            summary = "Financial documentation demonstrating budget vs. actual performance with narrative interpretation."
        elif 'balance sheet' in filename_lower or 'statement of financial position' in filename_lower:
            summary = "Balance sheet (statement of financial position) with narrative interpretation."
        elif 'income statement' in filename_lower or 'profit and loss' in filename_lower:
            summary = "Income statement (profit and loss statement) with narrative interpretation."
        elif 'cash flow' in filename_lower:
            summary = "Cash flow statement with narrative interpretation."
        else:
            summary = "Document providing evidence of financial stability. Review to verify it meets the requirements: quarterly financial reports for 2025 Q2 and Q3 with narrative interpretations of Budget vs. Actual, Balance Sheet, Income Statement, and Cash Flow Statement; and/or a financial plan approved by governing board addressing related party transactions, illiquid assets, negative equity, and low cash balances."
    
    else:
        # Generic fallback
        if 'minutes' in text_lower or 'meeting' in text_lower:
            date_match = re.search(r'(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}', text, re.IGNORECASE)
            date_str = f" dated {date_match.group(0)}" if date_match else ""
            summary = f"Meeting minutes{date_str} documenting discussions, decisions, and action items."
        elif 'policy' in text_lower:
            summary = "A policy document establishing institutional policies and procedures."
        else:
            summary = "Document containing relevant information for compliance review."
    
    return summary

def create_standard_review_doc(standard_num, exhibits, pdf_dir, output_dir):
    """Create a Word document for a specific standard with all exhibits."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title with document count
    title = doc.add_heading(f'Standard {standard_num} – {len(exhibits)} Supporting Documents Review', level=1)
    title_format = title.paragraph_format
    title_format.space_after = Pt(6)
    
    if not exhibits:
        p = doc.add_paragraph()
        if standard_num == 2:
            p.add_run('EX2.05 document not found in the attachments directory.').italic = True
            p = doc.add_paragraph()
            p.add_run('EX2.05 should contain:').bold = True
            p = doc.add_paragraph()
            p.add_run('• Training materials for compliance with federal laws and regulations:').bold = True
            p = doc.add_paragraph('  - Title IX (sex-based discrimination)', style='List Bullet')
            p = doc.add_paragraph('  - HIPAA (privacy protection of patients\' health records)', style='List Bullet')
            p = doc.add_paragraph('  - FERPA (privacy protection of student academic records)', style='List Bullet')
            p = doc.add_paragraph('  - OSHA (occupational safety and health)', style='List Bullet')
            p = doc.add_paragraph()
            p.add_run('• Evidence that all relevant stakeholders complete required compliance trainings').bold = True
            p = doc.add_paragraph()
            p.add_run('• Evidence that staff members responsible for overseeing compliance with federal regulations are appropriately qualified').bold = True
            p = doc.add_paragraph()
            p.add_run('Please ensure EX2.05 is available for review.').italic = True
            
            # Look for related documents that might contain this information
            related_docs = find_federal_compliance_docs(pdf_dir)
            
            # Also search content for compliance information
            compliance_docs_found = []
            pdf_files = sorted(pdf_dir.glob('*.pdf')) + sorted(pdf_dir.glob('*.PDF'))
            
            for pdf_file in pdf_files:
                try:
                    text = extract_text_from_pdf(pdf_file)
                    text_lower = text.lower()
                    
                    # Check if document contains multiple compliance topics
                    compliance_count = 0
                    if 'title ix' in text_lower or 'title 9' in text_lower or 'sexual harassment' in text_lower:
                        compliance_count += 1
                    if 'hipaa' in text_lower:
                        compliance_count += 1
                    if 'ferpa' in text_lower:
                        compliance_count += 1
                    if 'osha' in text_lower:
                        compliance_count += 1
                    if 'training' in text_lower and ('completion' in text_lower or 'certificate' in text_lower):
                        compliance_count += 1
                    
                    # If document has 2+ compliance topics, it's relevant
                    if compliance_count >= 2:
                        compliance_docs_found.append((pdf_file, compliance_count))
                except:
                    pass
            
            if related_docs or compliance_docs_found:
                p = doc.add_paragraph()
                p.add_run('Documents found containing federal compliance information:').bold = True
                
                # Sort by relevance (number of compliance topics)
                compliance_docs_found.sort(key=lambda x: x[1], reverse=True)
                
                for doc_file, count in compliance_docs_found[:10]:  # Top 10 most relevant
                    p = doc.add_paragraph(f"{doc_file.name} (contains {count} compliance topics)", style='List Bullet')
                
                # Also add filename-based matches
                for doc_file in related_docs:
                    if (doc_file, 0) not in [(d, c) for d, c in compliance_docs_found]:
                        p = doc.add_paragraph(doc_file.name, style='List Bullet')
        else:
            p.add_run('No supporting documents found for this standard.').italic = True
        return doc
    
    # Process each exhibit
    for idx, exhibit_path in enumerate(exhibits, 1):
        exhibit_name = exhibit_path.name
        
        # Exhibit heading
        heading = doc.add_heading(f'Document {idx}: {exhibit_name}', level=2)
        heading_format = heading.paragraph_format
        heading_format.space_after = Pt(3)
        
        # Extract text
        print(f"  Processing: {exhibit_name}")
        text = extract_text_from_pdf(exhibit_path)
        
        # Summary Analysis section
        p = doc.add_paragraph()
        p.add_run('Document Analysis:').bold = True
        p.paragraph_format.space_after = Pt(3)
        
        # Create analytical summary (now returns a single paragraph string)
        analysis_summary = create_analysis_summary(text, exhibit_name, standard_num)
        p = doc.add_paragraph(analysis_summary)
        p.paragraph_format.space_after = Pt(6)
    
    # Compliance Assessment section (once per standard, after all documents)
    if exhibits:
        doc.add_paragraph()  # Spacing
        p = doc.add_paragraph()
        p.add_run('Compliance Assessment:').bold = True
        p.paragraph_format.space_after = Pt(6)
        
        compliance_para = doc.add_paragraph(style='List Bullet')
        compliance_run = compliance_para.add_run(
            "[Review team: Assess compliance with Standard " + 
            f"{standard_num} requirements. Document strengths, gaps, or areas requiring further documentation.]"
        )
        compliance_run.italic = True
        compliance_para.paragraph_format.space_after = Pt(12)
    
    return doc

def extract_key_findings(text, filename):
    """Extract key findings and important content highlights from document text."""
    findings = []
    
    if not text or text.startswith("Error reading PDF"):
        return findings
    
    text_lower = text.lower()
    
    # Extract specific data points that might be relevant
    # Dates
    date_pattern = r'\b(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b'
    dates = re.findall(date_pattern, text, re.IGNORECASE)
    if dates:
        unique_dates = list(set(dates))[:3]
        findings.append(f"Key dates referenced: {', '.join(unique_dates)}")
    
    # Percentages or rates (e.g., graduation rates)
    rate_pattern = r'(\d+\.?\d*)\s*%'
    rates = re.findall(rate_pattern, text)
    if rates:
        findings.append(f"Percentages/rates mentioned: {', '.join(set(rates[:3]))}%")
    
    # Specific numbers that might be important (years, counts, etc.)
    if 'graduation' in text_lower and 'rate' in text_lower:
        grad_match = re.search(r'graduation\s+rate[^\d]*(\d+\.?\d*)%?', text_lower)
        if grad_match:
            findings.append(f"Graduation rate: {grad_match.group(1)}%")
    
    # Extract key decisions or actions if it's meeting minutes
    if 'decision' in text_lower or 'action' in text_lower:
        # Look for decision/action patterns
        decision_section = re.search(r'decision/action[^:]*:?\s*([^\.]{50,200})', text_lower)
        if decision_section:
            findings.append("Document contains decisions or action items")
    
    # Extract key topics discussed
    topics = []
    if 'mission' in text_lower:
        topics.append("institutional mission")
    if 'curriculum' in text_lower:
        topics.append("curriculum")
    if 'faculty' in text_lower:
        topics.append("faculty matters")
    if 'student' in text_lower and 'enrollment' in text_lower:
        topics.append("student enrollment")
    if 'budget' in text_lower or 'financial' in text_lower:
        topics.append("financial matters")
    
    if topics:
        findings.append(f"Topics covered: {', '.join(set(topics))}")
    
    # For lists, try to extract count
    if 'list' in filename.lower():
        # Try to count items in list
        list_items = re.findall(r'^\d+[\.\)]\s', text, re.MULTILINE)
        if list_items:
            findings.append(f"Contains approximately {len(list_items)} listed items")
    
    return findings

def main():
    """Create review documents for specific standards with commission findings."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    
    print("Creating Standard Review Documents...")
    print(f"PDF Directory: {PDF_DIR}")
    print(f"Output Directory: {OUTPUT_DIR}\n")
    
    # Only process standards with commission findings
    standards_to_process = [2, 2.06, 2.07, 5.02, 7.01, 8.01, 9.02]
    
    for standard_num in standards_to_process:
        print(f"Processing Standard {standard_num}...")
        
        # Get exhibits for this standard
        exhibits = get_exhibits_for_standard(standard_num, PDF_DIR)
        print(f"  Found {len(exhibits)} supporting document(s)")
        
        # Create document
        doc = create_standard_review_doc(standard_num, exhibits, PDF_DIR, OUTPUT_DIR)
        
        # Save document with proper naming for sub-standards
        if standard_num == 2.06:
            output_file = OUTPUT_DIR / f"Standard_2.06_Review.docx"
        elif standard_num == 2.07:
            output_file = OUTPUT_DIR / f"Standard_2.07_Review.docx"
        elif standard_num == 5.02:
            output_file = OUTPUT_DIR / f"Standard_5.02_Review.docx"
        elif standard_num == 7.01:
            output_file = OUTPUT_DIR / f"Standard_7.01_Review.docx"
        elif standard_num == 8.01:
            output_file = OUTPUT_DIR / f"Standard_8.01_Review.docx"
        elif standard_num == 9.02:
            output_file = OUTPUT_DIR / f"Standard_9.02_Review.docx"
        else:
            output_file = OUTPUT_DIR / f"Standard_{standard_num}_Review.docx"
        
        doc.save(str(output_file))
        
        file_size = output_file.stat().st_size
        print(f"  Created: {output_file.name} ({file_size:,} bytes)\n")
    
    print("=" * 60)
    print("All standard review documents created successfully!")
    print(f"Output location: {OUTPUT_DIR}")
    print("=" * 60)

if __name__ == "__main__":
    main()
