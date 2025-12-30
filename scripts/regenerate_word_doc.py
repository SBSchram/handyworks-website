"""
Regenerate Word Document from Existing Firebase Data
Only creates the Word document, doesn't touch Stripe or Firebase
"""

import firebase_admin
from firebase_admin import credentials, firestore
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def initialize_firebase():
    """Initialize Firebase Admin SDK"""
    try:
        app = firebase_admin.get_app()
        print("Using existing Firebase app")
    except ValueError:
        cred = credentials.Certificate('serviceAccountKey.json')
        firebase_admin.initialize_app(cred)
        print("Initialized new Firebase app")
    
    return firestore.client()

def extract_lastname(full_name):
    """Extract last name from full name"""
    if not full_name or full_name == 'N/A':
        return 'Customer'
    
    parts = full_name.strip().split()
    return parts[-1] if parts else 'Customer'

def get_all_invoices(db):
    """Retrieve all invoices with payment links"""
    print("\nRetrieving invoices from Firebase...")
    
    invoices_ref = db.collection('handyworks_invoices')
    invoices = invoices_ref.stream()
    
    invoice_data = []
    
    for invoice_doc in invoices:
        invoice = invoice_doc.to_dict()
        
        # Only include invoices with payment links
        if invoice.get('stripe_payment_link_url'):
            invoice_data.append({
                'customer_name': invoice.get('customer_name', 'N/A'),
                'payment_link': invoice.get('stripe_payment_link_url', ''),
            })
    
    # Sort by last name
    invoice_data.sort(key=lambda x: extract_lastname(x.get('customer_name', 'N/A')))
    
    print(f"Found {len(invoice_data)} invoices with payment links")
    return invoice_data

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '008080')
    rPr.append(c)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink

def create_correction_document(invoices, output_file):
    """Create Word document with all corrections - SINGLE SPACING"""
    print(f"\nCreating Word document with {len(invoices)} entries...")
    
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    for idx, invoice in enumerate(invoices, 1):
        lastname = extract_lastname(invoice['customer_name'])
        payment_url = invoice['payment_link']
        
        print(f"  Adding entry {idx}/{len(invoices)}: Dr. {lastname}")
        
        # Greeting
        p1 = doc.add_paragraph(f"Dear Dr. {lastname},")
        
        # Body text
        p2 = doc.add_paragraph("Our invoice had an incomplete payment link to Stripe.")
        
        # Payment link
        p3 = doc.add_paragraph()
        p3.add_run("Here is the correct link: ")
        add_hyperlink(p3, "Click here to pay via Stripe", payment_url)
        
        # Apology
        p4 = doc.add_paragraph("We apologize for the inconvenience.")
        
        # Signature
        p5 = doc.add_paragraph("Best regards,")
        p6 = doc.add_paragraph("Dr. Steve")
        
        # Separator line (except for last entry)
        if idx < len(invoices):
            doc.add_paragraph("_" * 80)
    
    doc.save(output_file)
    print(f"\n✅ Document saved: {output_file}")

def main():
    """Main function"""
    print("=" * 60)
    print("Regenerate Word Document (Single Spacing)")
    print("=" * 60)
    
    # Initialize Firebase
    db = initialize_firebase()
    
    # Get all invoices with payment links
    invoices = get_all_invoices(db)
    
    if not invoices:
        print("\n❌ No invoices with payment links found.")
        return
    
    # Create Word document
    output_file = 'All_Invoice_Corrections_PERMANENT.docx'
    create_correction_document(invoices, output_file)
    
    # Print summary
    print("\n" + "=" * 60)
    print("SUMMARY")
    print("=" * 60)
    print(f"Total entries: {len(invoices)}")
    print(f"Output file: {output_file}")
    print("\n✅ Word document regenerated with single spacing!")
    print("✅ Ready to copy/paste into email replies!")

if __name__ == '__main__':
    main()
