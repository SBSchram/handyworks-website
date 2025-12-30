"""
Create Permanent Stripe Payment Links and Update Firebase
Replaces temporary checkout sessions with permanent payment links
"""

import firebase_admin
from firebase_admin import credentials, firestore
import stripe
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Initialize Stripe
# Set your Stripe secret key as environment variable or hardcode here
STRIPE_SECRET_KEY = os.environ.get('STRIPE_SECRET_KEY', 'sk_live_YOUR_KEY_HERE')
stripe.api_key = STRIPE_SECRET_KEY

def initialize_firebase():
    """Initialize Firebase Admin SDK"""
    try:
        app = firebase_admin.get_app()
        print("Using existing Firebase app")
    except ValueError:
        cred = credentials.Certificate('serviceAccountKey.json')
        firebase_admin.initialize_app(cred)
        print("Initialized new Firebase app")
    
    return firestore.client()

def extract_lastname(full_name):
    """Extract last name from full name"""
    if not full_name or full_name == 'N/A':
        return 'Customer'
    
    parts = full_name.strip().split()
    return parts[-1] if parts else 'Customer'

def create_permanent_payment_link(invoice_data):
    """
    Create a permanent Stripe Payment Link (not a checkout session)
    These links never expire and can be reused
    """
    try:
        # First, create a product
        product = stripe.Product.create(
            name=f"HandyWorks {invoice_data['year']} Annual Maintenance - {invoice_data['customer_name']}",
            description=f"Annual maintenance fee for {invoice_data['customer_name']}",
            metadata={
                'acct_num': str(invoice_data['acct_num']),
                'customer_name': invoice_data['customer_name'],
                'year': str(invoice_data['year']),
            }
        )
        
        # Create a price for this product
        amount_in_cents = int(float(invoice_data['amount']) * 100)
        price = stripe.Price.create(
            product=product.id,
            unit_amount=amount_in_cents,
            currency='usd',
        )
        
        # Create the permanent payment link
        payment_link = stripe.PaymentLink.create(
            line_items=[{
                'price': price.id,
                'quantity': 1,
            }],
            after_completion={
                'type': 'redirect',
                'redirect': {
                    'url': 'https://handyworks.com/payment-success.html',
                },
            },
            metadata={
                'acct_num': str(invoice_data['acct_num']),
                'customer_name': invoice_data['customer_name'],
                'year': str(invoice_data['year']),
                'invoice_amount': str(invoice_data['amount']),
            },
            billing_address_collection='auto',
            phone_number_collection={'enabled': True},
        )
        
        return {
            'id': payment_link.id,
            'url': payment_link.url,
            'product_id': product.id,
            'price_id': price.id,
        }
        
    except Exception as e:
        print(f"  ❌ Error creating payment link: {e}")
        return None

def get_all_invoices(db):
    """Retrieve all invoices"""
    print("\nRetrieving invoices from Firebase...")
    
    invoices_ref = db.collection('handyworks_invoices')
    invoices = invoices_ref.stream()
    
    invoice_data = []
    
    for invoice_doc in invoices:
        invoice = invoice_doc.to_dict()
        invoice['firebase_id'] = invoice_doc.id
        invoice_data.append(invoice)
    
    # Sort by last name
    invoice_data.sort(key=lambda x: extract_lastname(x.get('customer_name', 'N/A')))
    
    print(f"Found {len(invoice_data)} invoices")
    return invoice_data

def update_invoice_with_payment_link(db, firebase_id, payment_link_data):
    """Update invoice in Firebase with new permanent payment link"""
    try:
        db.collection('handyworks_invoices').document(firebase_id).update({
            'stripe_payment_link_id': payment_link_data['id'],
            'stripe_payment_link_url': payment_link_data['url'],
            'stripe_product_id': payment_link_data['product_id'],
            'stripe_price_id': payment_link_data['price_id'],
            'updated_at': firestore.SERVER_TIMESTAMP,
        })
        return True
    except Exception as e:
        print(f"  ❌ Error updating Firebase: {e}")
        return False

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '008080')
    rPr.append(c)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink

def create_correction_document(invoices, output_file):
    """Create Word document with all corrections"""
    print(f"\nCreating Word document with {len(invoices)} entries...")
    
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    for idx, invoice in enumerate(invoices, 1):
        lastname = extract_lastname(invoice['customer_name'])
        payment_url = invoice['new_payment_link']
        
        print(f"  Adding entry {idx}/{len(invoices)}: Dr. {lastname}")
        
        # Greeting
        p1 = doc.add_paragraph(f"Dear Dr. {lastname},")
        
        # Body text
        p2 = doc.add_paragraph("Our invoice had an incomplete payment link to Stripe.")
        
        # Payment link
        p3 = doc.add_paragraph()
        p3.add_run("Here is the correct link: ")
        add_hyperlink(p3, "Click here to pay via Stripe", payment_url)
        
        # Apology
        p4 = doc.add_paragraph("We apologize for the inconvenience.")
        
        # Signature
        p5 = doc.add_paragraph("Best regards,")
        p6 = doc.add_paragraph("Dr. Steve")
        
        # Separator line (except for last entry)
        if idx < len(invoices):
            doc.add_paragraph("_" * 80)
    
    doc.save(output_file)
    print(f"\n✅ Document saved: {output_file}")

def main():
    """Main function"""
    print("=" * 60)
    print("Create Permanent Stripe Payment Links")
    print("=" * 60)
    
    # Check Stripe key
    if STRIPE_SECRET_KEY == 'sk_live_YOUR_KEY_HERE':
        print("\n❌ ERROR: Please set your Stripe secret key!")
        print("Set environment variable: STRIPE_SECRET_KEY=sk_live_...")
        print("Or edit the script and replace 'sk_live_YOUR_KEY_HERE'")
        return
    
    # Initialize Firebase
    db = initialize_firebase()
    
    # Get all invoices
    invoices = get_all_invoices(db)
    
    if not invoices:
        print("\n❌ No invoices found.")
        return
    
    print(f"\n{'=' * 60}")
    print(f"Processing {len(invoices)} invoices...")
    print(f"{'=' * 60}\n")
    
    successful_updates = []
    failed_updates = []
    
    for idx, invoice in enumerate(invoices, 1):
        lastname = extract_lastname(invoice.get('customer_name', 'N/A'))
        print(f"[{idx}/{len(invoices)}] Processing Dr. {lastname}...")
        
        # Create permanent payment link
        payment_link_data = create_permanent_payment_link({
            'acct_num': invoice.get('acct_num', 'N/A'),
            'customer_name': invoice.get('customer_name', 'N/A'),
            'year': invoice.get('year', 'N/A'),
            'amount': invoice.get('amount', 0),
        })
        
        if not payment_link_data:
            failed_updates.append(invoice)
            continue
        
        print(f"  ✅ Created payment link: {payment_link_data['url'][:50]}...")
        
        # Update Firebase
        success = update_invoice_with_payment_link(
            db,
            invoice['firebase_id'],
            payment_link_data
        )
        
        if success:
            print(f"  ✅ Updated Firebase")
            invoice['new_payment_link'] = payment_link_data['url']
            successful_updates.append(invoice)
        else:
            failed_updates.append(invoice)
        
        print()
    
    # Create Word document with successful updates
    if successful_updates:
        output_file = 'All_Invoice_Corrections_PERMANENT.docx'
        create_correction_document(successful_updates, output_file)
    
    # Print summary
    print("\n" + "=" * 60)
    print("SUMMARY")
    print("=" * 60)
    print(f"Total invoices: {len(invoices)}")
    print(f"✅ Successfully updated: {len(successful_updates)}")
    print(f"❌ Failed: {len(failed_updates)}")
    
    if failed_updates:
        print("\nFailed invoices:")
        for inv in failed_updates:
            print(f"  - {inv.get('customer_name', 'N/A')} (Acct: {inv.get('acct_num', 'N/A')})")
    
    if successful_updates:
        print(f"\n✅ Word document created: All_Invoice_Corrections_PERMANENT.docx")
        print("✅ Firebase updated with permanent payment links")
        print("\n🎉 All done! The new links never expire and can be reused.")

if __name__ == '__main__':
    main()
