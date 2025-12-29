import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def markdown_to_docx(md_file_path, docx_file_path):
    """Convert a markdown file to a Word document."""
    # Read markdown file
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Create Word document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    lines = md_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            i += 1
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
            i += 1
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
            i += 1
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
            i += 1
        # Horizontal rule
        elif line.startswith('---'):
            p = doc.add_paragraph()
            p.add_run('─' * 50).bold = True
            i += 1
        # Bold text (simple pattern)
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
            i += 1
        # Bullet points
        elif line.strip().startswith('- '):
            text = line.strip()[2:]
            # Handle bold in bullet points
            if '**' in text:
                p = doc.add_paragraph(style='List Bullet')
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)
            else:
                doc.add_paragraph(text, style='List Bullet')
            i += 1
        # Numbered lists
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            doc.add_paragraph(text, style='List Number')
            i += 1
        # Regular paragraph
        else:
            # Handle inline formatting
            p = doc.add_paragraph()
            # Simple markdown parsing for bold
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
            i += 1
    
    # Save document
    doc.save(docx_file_path)
    print(f"Converted: {md_file_path.name} -> {docx_file_path.name}")

def convert_all_md_files(workspace_dir):
    """Convert all markdown files in the workspace to Word documents."""
    workspace_path = Path(workspace_dir)
    
    # Find all .md files
    md_files = list(workspace_path.glob('*.md'))
    
    if not md_files:
        print("No markdown files found in the workspace.")
        return
    
    print(f"Found {len(md_files)} markdown file(s) to convert...")
    
    for md_file in md_files:
        # Skip if it's in a subdirectory (like scripts folder)
        if md_file.parent != workspace_path:
            continue
            
        docx_file = md_file.with_suffix('.docx')
        try:
            markdown_to_docx(md_file, docx_file)
        except Exception as e:
            print(f"Error converting {md_file.name}: {str(e)}")

if __name__ == "__main__":
    workspace = r"C:\Users\sbsch\Documents\handyworks-website"
    convert_all_md_files(workspace)
    print("\nConversion complete!")
