"""
Generate a Word document version of the evidence analysis summary for commissioners.
This script creates a clean, actionable summary showing evidence present, gaps, and recommendations.
"""
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import from evidence_analysis_framework module
try:
    from evidence_analysis_framework import (
        COMMISSION_FINDINGS,
        PDF_DIR,
        extract_text_from_pdf,
        analyze_standard_evidence
    )
except ImportError as e:
    print("ERROR: Could not import evidence_analysis_framework module.")
    print("Please ensure evidence_analysis_framework.py is in the same directory.")
    print(f"Import error: {e}")
    sys.exit(1)

def generate_commissioner_summary_docx(output_path, pdf_dir=None, commission_findings=None, school_name="Institution"):
    """
    Generate Word document with clean, actionable summary for commissioners.
    
    Args:
        output_path: Path to save the Word document
        pdf_dir: Path to directory containing PDFs (defaults to PDF_DIR)
        commission_findings: Dict of findings (defaults to COMMISSION_FINDINGS)
        school_name: Name of the school/institution
    """
    if pdf_dir is None:
        pdf_dir = PDF_DIR
    
    if commission_findings is None:
        commission_findings = COMMISSION_FINDINGS
    
    # Collect all PDF files
    pdf_files = set()
    for f in pdf_dir.glob('*.pdf'):
        pdf_files.add(f)
    for f in pdf_dir.glob('*.PDF'):
        pdf_files.add(f)
    pdf_files = sorted(pdf_files)
    
    # Import the flexible document finder
    from evidence_analysis_framework import find_documents_for_standard
    
    # Create Word document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Evidence Analysis Summary for Commission Review', level=1)
    title_format = title.paragraph_format
    title_format.space_after = Pt(12)
    
    subtitle = doc.add_heading(f'{school_name} - Supplemental Report', level=2)
    subtitle_format = subtitle.paragraph_format
    title_format.space_after = Pt(12)
    
    # Purpose statement
    p = doc.add_paragraph()
    p.add_run('Purpose: ').bold = True
    p.add_run('Identify evidence present and gaps to assist commission decision-making.')
    
    p = doc.add_paragraph()
    p.add_run('This analysis does not make compliance determinations.').italic = True
    p.paragraph_format.space_after = Pt(12)
    
    # Analyze each standard
    for standard_num in commission_findings.keys():
        # Find documents using flexible pattern matching
        relevant_docs = find_documents_for_standard(standard_num, pdf_files)
        
        # Standard heading
        heading = doc.add_heading(f'Standard {standard_num}', level=1)
        heading_format = heading.paragraph_format
        heading_format.space_after = Pt(12)
        
        finding_info = commission_findings[standard_num]
        
        # Finding
        p = doc.add_paragraph()
        p.add_run('Commission Finding: ').bold = True
        p.add_run(finding_info['finding'])
        p.paragraph_format.space_after = Pt(6)
        
        # Documents provided
        p = doc.add_paragraph()
        p.add_run(f'Documents Provided: {len(relevant_docs)}').bold = True
        p.paragraph_format.space_after = Pt(3)
        
        if not relevant_docs:
            p = doc.add_paragraph('GAP: All required evidence appears to be missing.')
            p.paragraph_format.space_after = Pt(12)
            continue
        
        # List documents
        for doc_file in relevant_docs:
            doc.add_paragraph(doc_file.name, style='List Bullet')
        
        # Analyze evidence (pass findings as parameter)
        analysis = analyze_standard_evidence(standard_num, relevant_docs, commission_findings)
        
        # Evidence present
        p = doc.add_paragraph()
        evidence_count = len(analysis['summary']['evidence_present'])
        total_count = len(analysis['evidence_by_requirement'])
        p.add_run(f'Evidence Present: {evidence_count} of {total_count} required items').bold = True
        p.paragraph_format.space_after = Pt(6)
        
        if analysis['summary']['evidence_present']:
            for req_key in analysis['summary']['evidence_present']:
                req_info = analysis['evidence_by_requirement'][req_key]
                docs_list = [d['document'] for d in req_info['supporting_documents']]
                
                p = doc.add_paragraph()
                run = p.add_run('[PRESENT] ')
                run.bold = True
                p.add_run(req_info['description'])
                
                if docs_list:
                    doc_text = ', '.join(docs_list[:2])
                    if len(docs_list) > 2:
                        doc_text += '...'
                    p2 = doc.add_paragraph(f'  Document(s): {doc_text}', style='List Bullet 2')
        else:
            p = doc.add_paragraph('None identified')
        
        # Gaps
        p = doc.add_paragraph()
        gaps_count = len(analysis['summary']['gaps'])
        p.add_run(f'Gaps Identified: {gaps_count}').bold = True
        p.paragraph_format.space_after = Pt(6)
        
        if analysis['summary']['gaps']:
            for gap in analysis['summary']['gaps']:
                p = doc.add_paragraph()
                run = p.add_run('[MISSING] ')
                run.bold = True
                p.add_run(gap['description'])
        else:
            p = doc.add_paragraph('None identified - all required evidence appears present in provided documents')
        
        # Unreadable documents
        unreadable = [d for d in analysis['document_analyses'] if not d['readable']]
        if unreadable:
            p = doc.add_paragraph()
            p.add_run(f'Documents Requiring Physical Review: {len(unreadable)}').bold = True
            p.paragraph_format.space_after = Pt(6)
            for doc_analysis in unreadable:
                doc.add_paragraph(f"{doc_analysis['filename']} (unreadable)", style='List Bullet')
        
        # Spacing between standards
        doc.add_paragraph()
    
    # Analysis notes
    heading = doc.add_heading('Analysis Notes for Commissioners', level=1)
    heading_format = heading.paragraph_format
    heading_format.space_after = Pt(12)
    
    notes = [
        "This analysis is based on automated text extraction and keyword matching.",
        "Commissioners should review actual documents to verify all findings.",
        "Confidence levels indicate strength of automated match: HIGH (strong match), MEDIUM (good match), LOW (partial match - manual review recommended).",
        "Documents marked 'unreadable' require physical review.",
        "Gaps indicate missing evidence - may require additional documentation from institution."
    ]
    
    for i, note in enumerate(notes, 1):
        doc.add_paragraph(f"{i}. {note}", style='List Number')
    
    # Save document
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    
    file_size = output_path.stat().st_size
    print(f"Evidence Analysis Summary saved to: {output_path}")
    print(f"File size: {file_size:,} bytes")
    return output_path

if __name__ == "__main__":
    from datetime import datetime
    
    output_dir = Path(__file__).parent.parent
    output_file = output_dir / "VU_Evidence_Analysis_Summary.docx"
    
    print("="*80)
    print("Generating Evidence Analysis Summary for Vitality University")
    print("="*80)
    print()
    
    # Use default COMMISSION_FINDINGS and PDF_DIR from framework
    generate_commissioner_summary_docx(output_file)
    
    print()
    print("="*80)
    print("Analysis complete!")
    print("="*80)
