import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import sys

def markdown_to_docx(md_file_path, docx_file_path):
    """Convert a markdown file to a Word document with improved formatting."""
    try:
        # Read markdown file
        with open(md_file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Create Word document
        doc = Document()
        
        # Set default font and style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Set paragraph spacing
        paragraph_format = style.paragraph_format
        paragraph_format.space_after = Pt(6)
        
        lines = md_content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].rstrip()
            
            # Skip empty lines (but add minimal spacing)
            if not line:
                i += 1
                if i < len(lines) and lines[i].strip():  # Only add space if next line has content
                    doc.add_paragraph()
                continue
            
            # Headers
            if line.startswith('# '):
                heading = doc.add_heading(line[2:], level=1)
                heading_format = heading.paragraph_format
                heading_format.space_after = Pt(12)
                i += 1
            elif line.startswith('## '):
                heading = doc.add_heading(line[3:], level=2)
                heading_format = heading.paragraph_format
                heading_format.space_after = Pt(10)
                i += 1
            elif line.startswith('### '):
                heading = doc.add_heading(line[4:], level=3)
                heading_format = heading.paragraph_format
                heading_format.space_after = Pt(8)
                i += 1
            elif line.startswith('#### '):
                heading = doc.add_heading(line[5:], level=4)
                heading_format = heading.paragraph_format
                heading_format.space_after = Pt(6)
                i += 1
            # Horizontal rule
            elif line.startswith('---'):
                p = doc.add_paragraph()
                p.add_run('─' * 50).bold = True
                p.paragraph_format.space_after = Pt(12)
                i += 1
            # Bold text (simple pattern)
            elif line.strip().startswith('**') and line.strip().endswith('**') and len(line.strip()) > 4:
                p = doc.add_paragraph()
                run = p.add_run(line.strip()[2:-2])
                run.bold = True
                i += 1
            # Bullet points
            elif line.strip().startswith('- '):
                text = line.strip()[2:]
                p = doc.add_paragraph(style='List Bullet')
                # Handle bold in bullet points
                if '**' in text:
                    parts = re.split(r'(\*\*.*?\*\*)', text)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part[2:-2])
                            run.bold = True
                        else:
                            p.add_run(part)
                else:
                    p.add_run(text)
                i += 1
            # Numbered lists
            elif re.match(r'^\d+\.\s', line.strip()):
                text = re.sub(r'^\d+\.\s', '', line.strip())
                p = doc.add_paragraph(style='List Number')
                if '**' in text:
                    parts = re.split(r'(\*\*.*?\*\*)', text)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part[2:-2])
                            run.bold = True
                        else:
                            p.add_run(part)
                else:
                    p.add_run(text)
                i += 1
            # Regular paragraph
            else:
                p = doc.add_paragraph()
                # Handle inline formatting (bold)
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)
                i += 1
        
        # Validate document before saving
        if len(doc.paragraphs) == 0:
            raise ValueError("Document has no content")
        
        # Save document with error handling
        try:
            docx_file_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(docx_file_path))
            
            # Verify file was created and has content
            if not docx_file_path.exists():
                raise IOError(f"File was not created: {docx_file_path}")
            
            file_size = docx_file_path.stat().st_size
            if file_size < 1000:  # Word docs should be at least 1KB
                raise ValueError(f"File is too small ({file_size} bytes), may be corrupted")
            
            print(f"[OK] Converted: {md_file_path.name} -> {docx_file_path.name} ({file_size:,} bytes)")
            return True
            
        except PermissionError:
            print(f"[ERROR] Permission denied: Cannot write to {docx_file_path}")
            return False
        except Exception as e:
            print(f"[ERROR] Error saving {docx_file_path.name}: {str(e)}")
            return False
            
    except Exception as e:
        print(f"[ERROR] Error converting {md_file_path.name}: {str(e)}")
        return False

def fix_temp_environment():
    """Check and fix temp environment variables."""
    temp_dir = Path(os.environ.get('TEMP', os.environ.get('TMP', '')))
    
    if not temp_dir:
        print("Warning: TEMP environment variable not set")
        return False
    
    if not temp_dir.exists():
        print(f"Warning: Temp directory does not exist: {temp_dir}")
        try:
            temp_dir.mkdir(parents=True, exist_ok=True)
            print(f"Created temp directory: {temp_dir}")
        except Exception as e:
            print(f"Error creating temp directory: {e}")
            return False
    
    # Test write permissions
    try:
        test_file = temp_dir / "test_write_permissions.tmp"
        test_file.write_text("test")
        test_file.unlink()
        print(f"[OK] Temp directory is writable: {temp_dir}")
        return True
    except Exception as e:
        print(f"[WARNING] Temp directory is not writable: {temp_dir} - {e}")
        return False

def convert_all_md_files(workspace_dir, output_dir=None):
    """Convert all markdown files in the workspace to Word documents."""
    workspace_path = Path(workspace_dir)
    
    if output_dir:
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
    else:
        output_path = workspace_path
    
    # Find all .md files in workspace root
    md_files = [f for f in workspace_path.glob('*.md') if f.parent == workspace_path]
    
    if not md_files:
        print("No markdown files found in the workspace root.")
        return
    
    print(f"\nFound {len(md_files)} markdown file(s) to convert...")
    print(f"Output directory: {output_path}\n")
    
    # Check temp environment
    print("Checking temp directory...")
    fix_temp_environment()
    print()
    
    success_count = 0
    for md_file in md_files:
        docx_file = output_path / md_file.with_suffix('.docx').name
        if markdown_to_docx(md_file, docx_file):
            success_count += 1
    
    print(f"\n{'='*60}")
    print(f"Conversion complete: {success_count}/{len(md_files)} files converted successfully")
    print(f"{'='*60}")

if __name__ == "__main__":
    workspace = r"C:\Users\sbsch\Documents\handyworks-website"
    
    # Optionally specify a different output directory
    # output = r"C:\Users\sbsch\Documents\handyworks-website\Word_Documents"
    
    convert_all_md_files(workspace)
