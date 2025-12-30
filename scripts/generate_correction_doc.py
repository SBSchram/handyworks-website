"""
Generate Single Word Document with All Invoice Corrections
Creates one document with all 67 customer entries pre-filled
"""

import firebase_admin
from firebase_admin import credentials, firestore
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def initialize_firebase():
    """Initialize Firebase Admin SDK"""
    try:
        app = firebase_admin.get_app()
        print("Using existing Firebase app")
    except ValueError:
        # Initialize new app - UPDATE THIS PATH
        cred = credentials.Certificate('serviceAccountKey.json')
        firebase_admin.initialize_app(cred)
        print("Initialized new Firebase app")
    
    return firestore.client()

def add_hyperlink(paragraph, text, url):
    """
    Add a hyperlink to a paragraph
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Set the run's style to hyperlink style
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    
    # Set color
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '008080')  # Teal color
    rPr.append(c)
    
    # Set underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)

    return hyperlink

def extract_lastname(full_name):
    """Extract last name from full name"""
    if not full_name or full_name == 'N/A':
        return 'Customer'
    
    parts = full_name.strip().split()
    return parts[-1] if parts else 'Customer'

def get_all_invoices_with_payment_links(db):
    """Retrieve all invoices that have Stripe payment links"""
    print("\nRetrieving invoices from Firebase...")
    
    invoices_ref = db.collection('handyworks_invoices')
    invoices = invoices_ref.stream()
    
    invoice_data = []
    
    for invoice_doc in invoices:
        invoice = invoice_doc.to_dict()
        
        # Only include invoices with payment links
        if invoice.get('stripe_payment_link_url'):
            invoice_data.append({
                'id': invoice_doc.id,
                'customer_name': invoice.get('customer_name', 'N/A'),
                'lastname': extract_lastname(invoice.get('customer_name', 'N/A')),
                'payment_link': invoice.get('stripe_payment_link_url', ''),
                'acct_num': invoice.get('acct_num', 'N/A'),
                'year': invoice.get('year', 'N/A'),
            })
    
    # Sort by last name for easier navigation
    invoice_data.sort(key=lambda x: x['lastname'])
    
    print(f"Found {len(invoice_data)} invoices with payment links")
    return invoice_data

def create_correction_document(invoices, output_file):
    """Create a single Word document with all corrections"""
    print(f"\nCreating Word document with {len(invoices)} entries...")
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Add each customer entry
    for idx, invoice in enumerate(invoices, 1):
        lastname = invoice['lastname']
        payment_url = invoice['payment_link']
        
        print(f"  Adding entry {idx}/{len(invoices)}: Dr. {lastname}")
        
        # Greeting
        p1 = doc.add_paragraph()
        p1.add_run(f"Dear Dr. {lastname},")
        
        # Blank line
        doc.add_paragraph()
        
        # Body text
        doc.add_paragraph("Our invoice had an incomplete payment link to Stripe.")
        
        # Blank line
        doc.add_paragraph()
        
        # Payment link paragraph
        p2 = doc.add_paragraph()
        p2.add_run("Here is the correct link: ")
        add_hyperlink(p2, "Click here to pay via Stripe", payment_url)
        
        # Blank line
        doc.add_paragraph()
        
        # Apology
        doc.add_paragraph("We apologize for the inconvenience.")
        
        # Blank line
        doc.add_paragraph()
        
        # Signature
        doc.add_paragraph("Best regards,")
        doc.add_paragraph("Dr. Steve")
        
        # Separator line (except for last entry)
        if idx < len(invoices):
            doc.add_paragraph("_" * 80)
            doc.add_paragraph()  # Extra blank line between entries
    
    # Save document
    doc.save(output_file)
    print(f"\n✅ Document saved: {output_file}")

def main():
    """Main function"""
    print("=" * 60)
    print("Invoice Payment Link Correction Document Generator")
    print("=" * 60)
    
    # Initialize Firebase
    db = initialize_firebase()
    
    # Get all invoices with payment links
    invoices = get_all_invoices_with_payment_links(db)
    
    if not invoices:
        print("\n❌ No invoices with payment links found.")
        return
    
    # Create the Word document
    output_file = 'All_Invoice_Corrections.docx'
    create_correction_document(invoices, output_file)
    
    # Print summary
    print("\n" + "=" * 60)
    print("SUMMARY")
    print("=" * 60)
    print(f"Total entries: {len(invoices)}")
    print(f"Output file: {output_file}")
    print("\nEntries are sorted alphabetically by last name.")
    print("Each entry is separated by a line.")
    print("\n✅ Ready to copy/paste into email replies!")

if __name__ == '__main__':
    main()
