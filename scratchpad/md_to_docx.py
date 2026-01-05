#!/usr/bin/env python3
"""
Robust Markdown to Word Document Converter

Converts Markdown (.md) files to Word (.docx) documents with proper formatting.
Handles headings, lists, bold, italic, links, code blocks, and more.

Features:
    - Headings (H1-H6)
    - Unordered and ordered lists (with nested support)
    - Bold (**text** or __text__)
    - Italic (*text* or _text_)
    - Inline code (`code`)
    - Code blocks (```code```)
    - Links ([text](url))
    - Blockquotes (>)
    - Horizontal rules (---)

Usage:
    python scripts/md_to_docx.py input.md [output.docx]
    
    If output is not specified, creates input.docx in the same directory as input.md

Examples:
    python scripts/md_to_docx.py README.md
    python scripts/md_to_docx.py document.md output.docx
    python scripts/md_to_docx.py ../notes.md ./converted.docx

Requirements:
    python-docx library (install with: pip install python-docx)
"""

import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Style for hyperlinks (blue, underlined)
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink

def parse_inline_formatting(text, paragraph):
    """Parse inline markdown formatting (bold, italic, code, links) using regex"""
    if not text:
        return
    
    # Process text by finding all markdown elements and their positions
    # Order: links, code, bold, italic (to avoid conflicts)
    pos = 0
    text_len = len(text)
    
    while pos < text_len:
        # Find the next markdown element
        next_link = re.search(r'\[([^\]]+)\]\(([^\)]+)\)', text[pos:])
        next_code = re.search(r'`([^`]+)`', text[pos:])
        next_bold = re.search(r'\*\*([^*]+)\*\*|__([^_]+)__', text[pos:])
        next_italic = re.search(r'(?<!\*)\*([^*]+)\*|(?<!_)_([^_]+)_', text[pos:])
        
        # Find the earliest match
        matches = []
        if next_link:
            matches.append(('link', next_link.start() + pos, next_link.end() + pos, next_link))
        if next_code:
            matches.append(('code', next_code.start() + pos, next_code.end() + pos, next_code))
        if next_bold:
            matches.append(('bold', next_bold.start() + pos, next_bold.end() + pos, next_bold))
        if next_italic:
            matches.append(('italic', next_italic.start() + pos, next_italic.end() + pos, next_italic))
        
        if not matches:
            # No more formatting, add rest as plain text
            paragraph.add_run(text[pos:])
            break
        
        # Sort by position and process earliest
        matches.sort(key=lambda x: x[1])
        match_type, match_start, match_end, match_obj = matches[0]
        
        # Add text before the match
        if match_start > pos:
            process_text_formatting(text[pos:match_start], paragraph)
        
        # Process the match
        if match_type == 'link':
            add_hyperlink(paragraph, match_obj.group(2), match_obj.group(1))
        elif match_type == 'code':
            run = paragraph.add_run(match_obj.group(1))
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        elif match_type == 'bold':
            bold_text = match_obj.group(1) or match_obj.group(2)
            paragraph.add_run(bold_text).bold = True
        elif match_type == 'italic':
            italic_text = match_obj.group(1) or match_obj.group(2)
            paragraph.add_run(italic_text).italic = True
        
        pos = match_end

def process_text_formatting(text, paragraph):
    """Process plain text (no links, code, or complex formatting)"""
    if not text:
        return
    
    # Simple approach: just add the text as-is
    # (Complex formatting like nested bold/italic is handled by parse_inline_formatting)
    paragraph.add_run(text)

def convert_markdown_to_docx(md_file_path, docx_file_path=None):
    """
    Convert a Markdown file to a Word document.
    
    Args:
        md_file_path: Path to input .md file
        docx_file_path: Path to output .docx file (optional)
    
    Returns:
        Path to created .docx file
    """
    md_path = Path(md_file_path)
    
    if not md_path.exists():
        raise FileNotFoundError(f"Markdown file not found: {md_file_path}")
    
    # Determine output path
    if docx_file_path:
        output_path = Path(docx_file_path)
    else:
        output_path = md_path.with_suffix('.docx')
    
    # Read markdown file
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # Create Word document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    i = 0
    in_code_block = False
    code_block_language = None
    code_block_lines = []
    
    while i < len(lines):
        line = lines[i].rstrip('\n\r')
        original_line = line
        
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block - add as formatted code paragraph
                if code_block_lines:
                    code_para = doc.add_paragraph()
                    code_run = code_para.add_run('\n'.join(code_block_lines))
                    code_run.font.name = 'Courier New'
                    code_run.font.size = Pt(9)
                    code_para.style = 'No Spacing'
                code_block_lines = []
                in_code_block = False
                code_block_language = None
            else:
                # Start of code block
                in_code_block = True
                code_block_language = line[3:].strip() if len(line) > 3 else None
            i += 1
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # Skip empty lines (but preserve them between paragraphs)
        if not line.strip():
            # Only add blank line if previous line wasn't empty
            if i > 0 and lines[i-1].strip():
                doc.add_paragraph()
            i += 1
            continue
        
        # Headings (# ## ### etc.)
        if line.startswith('#'):
            level = 0
            while level < len(line) and line[level] == '#':
                level += 1
            
            heading_text = line[level:].strip()
            if heading_text:
                if level <= 6:
                    doc.add_heading(heading_text, level=min(level, 6))
                else:
                    # Level 7+ as bold paragraph
                    para = doc.add_paragraph()
                    para.add_run(heading_text).bold = True
            i += 1
            continue
        
        # Horizontal rule (--- or ***)
        if re.match(r'^[-*_]{3,}$', line.strip()):
            para = doc.add_paragraph()
            para.add_run('_' * 50)
            i += 1
            continue
        
        # Unordered lists (- or *)
        if re.match(r'^[\s]*[-*]\s+', line):
            indent_level = len(line) - len(line.lstrip())
            list_marker = re.match(r'^[\s]*[-*]\s+', line)
            list_text = line[list_marker.end():]
            
            # Determine list style based on indent
            if indent_level == 0:
                list_style = 'List Bullet'
            elif indent_level <= 4:
                list_style = 'List Bullet 2'
            else:
                list_style = 'List Bullet 3'
            
            para = doc.add_paragraph(style=list_style)
            parse_inline_formatting(list_text, para)
            i += 1
            continue
        
        # Ordered lists (1. 2. etc.)
        if re.match(r'^[\s]*\d+\.\s+', line):
            indent_level = len(line) - len(line.lstrip())
            list_marker = re.match(r'^[\s]*\d+\.\s+', line)
            list_text = line[list_marker.end():]
            
            # Determine list style based on indent
            if indent_level == 0:
                list_style = 'List Number'
            elif indent_level <= 4:
                list_style = 'List Number 2'
            else:
                list_style = 'List Number 3'
            
            para = doc.add_paragraph(style=list_style)
            parse_inline_formatting(list_text, para)
            i += 1
            continue
        
        # Blockquotes (>)
        if line.strip().startswith('>'):
            quote_text = line.lstrip('>').strip()
            para = doc.add_paragraph(style='Quote')
            parse_inline_formatting(quote_text, para)
            i += 1
            continue
        
        # Regular paragraph
        para = doc.add_paragraph()
        parse_inline_formatting(line, para)
        i += 1
    
    # Handle trailing code block
    if in_code_block and code_block_lines:
        code_para = doc.add_paragraph()
        code_run = code_para.add_run('\n'.join(code_block_lines))
        code_run.font.name = 'Courier New'
        code_run.font.size = Pt(9)
        code_para.style = 'No Spacing'
    
    # Save document
    doc.save(str(output_path))
    return output_path

def main():
    """Main function - command line interface"""
    if len(sys.argv) < 2:
        print("Usage: python scripts/md_to_docx.py <input.md> [output.docx]")
        print("\nExamples:")
        print("  python scripts/md_to_docx.py document.md")
        print("  python scripts/md_to_docx.py document.md output.docx")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    try:
        output_path = convert_markdown_to_docx(input_file, output_file)
        print(f"Successfully created: {output_path}")
    except FileNotFoundError as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"Error converting file: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == '__main__':
    main()
